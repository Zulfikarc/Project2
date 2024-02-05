import requests
from bs4 import BeautifulSoup
from docx import Document
from collections import Counter
import string


def get_latest_python_articles():
    url="https://www.python.org/"
    response=requests.get(url)

    if response.status_code==200:
        
        soup=BeautifulSoup(response.text,"html.parser")
        latest_articles=[]

        for article in soup.select(".blog-widget li"):
            title=article.a.text.strip()
            latest_articles.append(title)

        return latest_articles
    
    else:
        print(f"fail to retrieve information.status code:{response.status_code}")
        return []
    
if __name__=="__main__":
    python_articles=get_latest_python_articles()
    
    if python_articles:
        print("New News in the python.org section")
        for index,article in enumerate(python_articles,1):
            print(f"{index}. {article}")

    else:
        print("no article found")

    if python_articles:
        def create_word_document(articles, document_path):
            document = Document()
            document.add_heading('New News in the python.org section', level=1)

            for index, article in enumerate(python_articles, 1):
                document.add_paragraph(f"{index}. {article}")

            document.save(r'C:\Zulfikar\Data_Science_Internship\Project2\python_articles.docx')
            print("Word document created successfully.")
    else:
        print("No articles found")

    def count_word_frequency(document_path):
        try:
            # Load the Word document
            document = Document(document_path)

            # Extract text from paragraphs
            text = ""
            for paragraph in document.paragraphs:
                text += paragraph.text + " "

            # Remove punctuation and convert to lowercase
            text = text.translate(str.maketrans("", "", string.punctuation))
            words = text.lower().split()

            # Count word frequency
            word_frequency = Counter(words)

            return word_frequency

        except Exception as e:
            print(f"Error reading the document: {e}")
            return None

if __name__ == "__main__":
    python_articles = get_latest_python_articles()

    if python_articles:
        print("New News in the python.org section")
        for index, article in enumerate(python_articles, 1):
            print(f"{index}. {article}")

        create_word_document(python_articles, r'C:\Zulfikar\Data_Science_Internship\Project2\python_articles.docx')
        print("Word document created successfully.")

        word_frequency = count_word_frequency(r'C:\Zulfikar\Data_Science_Internship\Project2\python_articles.docx')

        if word_frequency:
            print("Word frequency count:")
            for word, count in word_frequency.items():
                print(f"{word}: {count}")
        else:
            print("Word frequency count failed.")
    else:
        print("No articles found")
#end project
        
   